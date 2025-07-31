# services/export_service.py

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Import des bibliothèques pour les différents formats
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch


class ExportService:
    """Service pour exporter les conversations dans différents formats"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def _build_markdown_content(self, chat_data: Dict) -> str:
        """
        Construit le contenu Markdown à partir des données de conversation
        
        Args:
            chat_data: Dictionnaire contenant 'summary' et 'history'
            
        Returns:
            Contenu formaté en Markdown
        """
        lines = []
        
        # Ajouter l'en-tête avec la date
        lines.append(f"# Conversation Exportée - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append("")
        
        # Ajouter le résumé du contexte s'il existe
        if 'summary' in chat_data and chat_data['summary']:
            lines.append("## Résumé du Contexte")
            lines.append("")
            lines.append(chat_data['summary'])
            lines.append("")
            lines.append("---")
            lines.append("")
        
        # Ajouter l'historique de la conversation
        lines.append("## Historique de la Conversation")
        lines.append("")
        
        if 'history' in chat_data and chat_data['history']:
            for i, message in enumerate(chat_data['history']):
                role = message.get('role', 'unknown')
                content = message.get('content', '')
                
                # Formater selon le rôle
                if role == 'user':
                    lines.append(f"### 👤 Utilisateur")
                elif role == 'assistant':
                    lines.append(f"### 🤖 Assistant")
                elif role == 'system':
                    lines.append(f"### ⚙️ Système")
                else:
                    lines.append(f"### {role.capitalize()}")
                
                lines.append("")
                
                # Ajouter le contenu avec indentation pour une meilleure lisibilité
                content_lines = content.split('\n')
                for line in content_lines:
                    if line.strip():
                        lines.append(f"> {line}")
                    else:
                        lines.append(">")
                
                lines.append("")
                lines.append("---")
                lines.append("")
        else:
            lines.append("*Aucun message dans l'historique*")
        
        return '\n'.join(lines)
    
    def _generate_md(self, markdown_content: str, output_path: str) -> None:
        """Génère un fichier Markdown"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
            
    def _generate_docx(self, markdown_content: str, output_path: str) -> None:
        """Génère un fichier Word (.docx)"""
        doc = Document()
        
        # Styles personnalisés
        title_style = doc.styles['Title']
        heading1_style = doc.styles['Heading 1']
        heading2_style = doc.styles['Heading 2']
        normal_style = doc.styles['Normal']
        
        # Parser le Markdown ligne par ligne
        lines = markdown_content.split('\n')
        in_quote = False
        quote_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if in_quote and quote_paragraph:
                    in_quote = False
                    quote_paragraph = None
                continue
                
            # Titre principal
            if line.startswith('# '):
                doc.add_paragraph(line[2:], style=title_style)
                
            # Sous-titres niveau 2
            elif line.startswith('## '):
                doc.add_paragraph(line[3:], style=heading1_style)
                
            # Sous-titres niveau 3
            elif line.startswith('### '):
                doc.add_paragraph(line[4:], style=heading2_style)
                
            # Citations (messages)
            elif line.startswith('> '):
                content = line[2:]
                if not in_quote:
                    quote_paragraph = doc.add_paragraph(content)
                    quote_paragraph.paragraph_format.left_indent = Inches(0.5)
                    in_quote = True
                else:
                    quote_paragraph.add_run('\n' + content)
                    
            # Ligne de séparation
            elif line == '---':
                doc.add_paragraph('_' * 50, style=normal_style)
                
            # Texte normal
            else:
                doc.add_paragraph(line, style=normal_style)
        
        # Sauvegarder le document
        doc.save(output_path)
        
    def _generate_pdf(self, markdown_content: str, output_path: str) -> None:
        """Génère un fichier PDF en utilisant ReportLab"""
        # Configuration du document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Conteneur pour les éléments du document
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30
        )
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=20
        )
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=12
        )
        normal_style = styles['Normal']
        quote_style = ParagraphStyle(
            'Quote',
            parent=styles['Normal'],
            leftIndent=36,
            rightIndent=36,
            textColor=colors.HexColor('#555555'),
            backColor=colors.HexColor('#f5f5f5')
        )
        
        # Parser le Markdown et construire le PDF
        lines = markdown_content.split('\n')
        quote_text = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Gérer les citations multi-lignes
            if line.startswith('> '):
                quote_text.append(line[2:])
                # Vérifier si c'est la dernière ligne de citation
                if i + 1 >= len(lines) or not lines[i + 1].strip().startswith('>'):
                    full_quote = '<br/>'.join(quote_text)
                    story.append(Paragraph(full_quote, quote_style))
                    story.append(Spacer(1, 12))
                    quote_text = []
                continue
                
            # Titre principal
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], title_style))
                
            # Sous-titres niveau 2
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], heading1_style))
                
            # Sous-titres niveau 3
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], heading2_style))
                
            # Ligne de séparation
            elif line == '---':
                story.append(Spacer(1, 12))
                # Créer une ligne horizontale
                data = [['―' * 80]]
                t = Table(data)
                t.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.grey),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 12))
                
            # Texte normal
            elif line:
                story.append(Paragraph(line, normal_style))
                story.append(Spacer(1, 6))
        
        # Construire le PDF
        doc.build(story)
        
    def generate_export(self, chat_data: Dict, output_path: str) -> Dict:
        """
        Génère l'export dans le format déterminé par l'extension du fichier
        
        Args:
            chat_data: Données de la conversation
            output_path: Chemin de sortie avec extension
            
        Returns:
            Dictionnaire avec le statut de l'opération
        """
        try:
            # Déterminer le format à partir de l'extension
            path = Path(output_path)
            extension = path.suffix.lower()
            
            # Construire le contenu Markdown (source commune)
            markdown_content = self._build_markdown_content(chat_data)
            
            # Générer selon le format
            if extension == '.md':
                self._generate_md(markdown_content, output_path)
            elif extension == '.docx':
                self._generate_docx(markdown_content, output_path)
            elif extension == '.pdf':
                self._generate_pdf(markdown_content, output_path)
            else:
                return {
                    'success': False,
                    'error': f'Format non supporté: {extension}'
                }
            
            # Vérifier que le fichier a été créé
            if not os.path.exists(output_path):
                return {
                    'success': False,
                    'error': 'Le fichier n\'a pas pu être créé'
                }
                
            return {
                'success': True,
                'path': output_path,
                'size': os.path.getsize(output_path)
            }
            
        except Exception as e:
            self.logger.error(f"Erreur lors de l'export: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }